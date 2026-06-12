from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── Title Page ──────────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PPE Compliance Detection System')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(26, 26, 46)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Challenges & Solutions Report')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(95, 99, 104)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Baseline YOLO11m vs SAM-Teacher (Grounding DINO + SAM 1 to YOLO11m)\n').font.size = Pt(12)
info.add_run('9,994 SAM-generated pseudo-labels | 55% mAP improvement (0.558 to 0.864)').font.size = Pt(11)

doc.add_page_break()

# ── Helpers ─────────────────────────────────────────────────────────────
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(26, 26, 46)
    return h

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bold_body(bold_text, normal_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)
    return p

def challenge_solution(challenge, cause, solution, outcome=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run('Challenge: ').bold = True
    p.add_run(challenge)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.add_run('Root Cause: ').bold = True
    p2.add_run(cause)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    p3.add_run('Solution: ').bold = True
    p3.add_run(solution)
    if outcome:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(10)
        p4.add_run('Outcome: ').bold = True
        p4.add_run(outcome)

# ═══════════════════════════════════════════════════════════════════════════════
# 1. Teacher Model Selection
# ═══════════════════════════════════════════════════════════════════════════════
heading('1. Teacher Model Selection', level=1)

heading('1.1 SAM 2 vs SAM 1: Dependency Hell', level=2)
challenge_solution(
    'SAM 2 (Segment Anything Model v2) refused to install reliably for automated pipelines.',
    'SAM 2 depends on Hydra config files that pip install does not reliably bundle. '
    'The package expects external YAML configuration files that are not part of the standard pip distribution. '
    'On Colab, importing SAM 2 consistently failed with missing config errors.',
    'Switched to SAM 1 (segment-anything package). SAM 1 ViT-H downloads as a single checkpoint file (2.4 GB) '
    'and works immediately after pip install, with no external configs or hidden dependencies.',
    'SAM 1 ViT-H produced 9,994 high-quality pseudo-labels across 1,263 images in 66 minutes on a Colab T4 GPU.'
)

heading('1.2 Grounding DINO vs Florence-2: Tokenizer Bug', level=2)
challenge_solution(
    'Florence-2 (Microsoft vision-language model) failed at inference due to a tokenizer configuration error.',
    'Newer versions of the Hugging Face transformers library introduced a breaking change: models with '
    'forced_bos_token_id in their config.json cause an assertion error during tokenizer initialization. '
    'Florence-2 pretrained config includes this parameter, making it incompatible with transformers >= 4.45.',
    'Switched to Grounding DINO (IDEA-Research/grounding-dino-base) which uses a standard BERT-based text encoder. '
    'It accepts natural language prompts and reliably detects all classes.',
    'Grounding DINO correctly identified all 6 target classes in the Construction-PPE dataset with box_threshold=0.30.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. Cross-Platform File Transfer
# ═══════════════════════════════════════════════════════════════════════════════
heading('2. Cross-Platform File Transfer', level=1)

heading('2.1 Zip Archives: Windows Backslashes Break on Linux', level=2)
challenge_solution(
    'Images zipped on Windows using PowerShell Compress-Archive would not extract properly on Google Colab (Linux).',
    'PowerShell Compress-Archive stores Windows-style backslash paths inside the ZIP. '
    'Linux tools like unzip interpret backslashes as escape characters, causing path resolution failures.',
    'Created a Python script using the zipfile module to create archives. '
    'Python zipfile normalizes all paths to forward slashes regardless of the host OS.',
    'All 1,416 training images transferred successfully to Colab with correct directory structure.'
)

heading('2.2 Dataset Download: The YOLO Train Trick', level=2)
challenge_solution(
    'The Construction-PPE dataset needed to be downloaded on Colab, but the Ultralytics dataset '
    'auto-download API only works during training, not as a standalone download.',
    'YOLO datasets on Roboflow/Ultralytics Hub are downloaded lazily: only when yolo train starts. '
    'There is no direct download endpoint for the raw images.',
    'Initiated a yolo train with epochs=1 as a download trigger, then immediately killed the training process. '
    'This populated the dataset cache with all images and labels.',
    'Dataset downloaded in ~2 minutes. Images moved to the labeler expected directory structure.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. Deployment Platform Selection
# ═══════════════════════════════════════════════════════════════════════════════
heading('3. Deployment Platform Selection', level=1)

heading('3.1 Free GPU Hosting: The Vanishing Landscape', level=2)
challenge_solution(
    'The dashboard needed GPU inference for live video processing, but every free GPU platform had a fatal limitation.',
    'Hugging Face Spaces GPU tier became Pro-only ($9/month) in 2025. Google Colab kills sessions after a few hours '
    'and cannot serve web traffic. Databricks Community Edition is CPU-only. Snowflake is a data warehouse, not '
    'an inference platform. Kaggle kernels have no persistent web serving.',
    'Adopted a two-tier architecture: Streamlit Cloud (free CPU) hosts the public dashboard; Modal (serverless GPU '
    'with $30/month free credits) handles on-demand inference. Modal only runs when a user clicks Process, costing '
    '$0.0001 per inference, effectively free forever for portfolio use.',
    'Dashboard always online at zero cost. GPU inference available on demand. Up to ~300,000 video processes '
    'before exhausting free credits.'
)

heading('3.2 Streamlit Cloud: Python Version & Repository Size', level=2)
challenge_solution(
    'Initial deployment to Streamlit Cloud failed with "Error installing requirements" on multiple attempts.',
    'Three issues combined: (1) Streamlit Cloud defaults to Python 3.14 which lacks complete package wheels. '
    '(2) The git repository contained 50 MB video files in history (373 MB total) causing clone timeout. '
    '(3) Ultralytics/torch dependencies were too heavy for the free 1GB RAM tier.',
    'Created a clean repository with only essential files (6 files, 200 KB). Set Python version to 3.11 '
    'via .python-version. Removed all video files and model weights from git. Used only streamlit, pandas, '
    'and plotly as dependencies (~50 MB install).',
    'Dashboard deploys in under 30 seconds. Always online with no memory pressure.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. Modal GPU Setup
# ═══════════════════════════════════════════════════════════════════════════════
heading('4. Modal GPU Setup: A Gauntlet of Errors', level=1)

heading('4.1 Image Build Failures', level=2)
challenge_solution(
    'Modal image builder failed repeatedly with different errors during package installation.',
    'Multiple cascading issues: (1) Torch download (888 MB) timed out on Modal internal PyPI mirror. '
    '(2) The debian_slim() function rejected the python="3.11" keyword argument (not supported in Modal '
    'client v1.5.0). (3) fastapi_endpoint decorator required FastAPI to be explicitly installed in the '
    'image (previously auto-included). (4) @modal.web_endpoint was deprecated and renamed.',
    'Switched torch installation from Modal mirror to PyTorch official CDN using run_commands(). '
    'Removed unsupported python kwarg. Added fastapi to pip_install list. Updated all decorators.',
    'Docker image builds in ~41 seconds. Torch + Ultralytics + OpenCV + FastAPI all available on T4 GPU.'
)

heading('4.2 Model Upload: Filesystem vs Cloud Mismatch', level=2)
challenge_solution(
    'Uploaded models were invisible to the inference function. The models/ directory existed locally '
    'but not in Modal cloud execution environment.',
    'modal run mounts only the Python script file, not subdirectories. Relative paths resolved to nothing '
    'in the cloud container. The modal.Mount API was not available in client v1.5.0.',
    'Created a standalone upload_models.py script using Modal Volume API directly: '
    'modal.Volume.from_name("ppe-models").batch_upload() to copy .pt files into persistent cloud storage. '
    'The inference function mounts the Volume at /models.',
    'Both 40.5 MB model files successfully persisted in Modal Volume and load correctly at inference time.'
)

heading('4.3 Video Encoding: The Broken Pipeline', level=2)
challenge_solution(
    'Annotated output videos failed to encode, producing corrupted files, empty files, then crashing '
    'the entire inference pipeline.',
    'A chain of failures: (1) OpenCV avc1 fourcc produced zero-byte files in Debian Docker. '
    '(2) Switching to mp4v produced files browsers could not play. (3) Attempting libx264 via ffmpeg '
    'failed because frame files had gaps in numbering; empty frames were not being saved. '
    '(4) Without frames, ffmpeg crashed and the output file was never created, causing FileNotFoundError.',
    'Simplified to a layered approach: (1) Send a single annotated JPEG frame as instant preview. '
    '(2) Encode video using OpenCV mp4v codec limited to 15 seconds for download. '
    '(3) Added per-model try/except so one model crashing does not lose the other results. '
    '(4) Added JPEG fallback if video encoding fails entirely.',
    'Preview displays instantly. Downloadable annotated videos available for both Baseline and SAM-Teacher. '
    'Dashboard shows results even if one model encounters an error.'
)

heading('4.4 Per-Model Error Isolation', level=2)
challenge_solution(
    'The SAM-Teacher model would crash silently, returning results for Baseline only. '
    'Users saw one-sided results with no error message.',
    'The processing loop processed both models sequentially without per-model error handling. '
    'If SAM failed on any frame, the entire function returned partial results or crashed entirely.',
    'Wrapped each model processing block in try/except. If SAM fails, results include '
    '{"sam": {"error": "..."}} while Baseline returns normally. Dashboard checks for error keys '
    'and displays warnings.',
    'Both models always return results independently. Robust to model-specific failures.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. Model Architecture
# ═══════════════════════════════════════════════════════════════════════════════
heading('5. Model Architecture Challenges', level=1)

heading('5.1 Class Remapping: Tensor Immutability', level=2)
challenge_solution(
    'SAM model outputs 6 classes (person + 5 PPE items), but the pipeline expects 11 classes '
    'matching the Construction-PPE dataset format.',
    'PyTorch inference tensors are immutable. Attempting to modify class IDs in-place raises '
    '"Inplace update to inference tensor" errors. Direct tensor manipulation is forbidden '
    'by PyTorch autograd engine.',
    'Remap class IDs on NumPy arrays after extraction: cls = preds.boxes.cls.cpu().numpy().astype(int), '
    'then apply remapping dict using NumPy operations. Negative classes (no_helmet, no_goggles, etc.) '
    'are inferred from the absence of corresponding positive detections near a person.',
    'Clean class remapping with no tensor mutation errors. All 11 expected classes represented in output.'
)

heading('5.2 Phase Isolation: Protecting Baseline Work', level=2)
challenge_solution(
    'Multi-phase ML projects risk accidentally overwriting or corrupting baseline results '
    'when iterating on improvements.',
    'Without explicit isolation rules, it is tempting to modify code in-place or import from '
    'sibling project folders, creating hidden dependencies. Models and weights are especially '
    'vulnerable to accidental overwrites.',
    'Enforced strict phase isolation: every phase gets its own folder (ppe-compliance/, ppe-sam-phase/, '
    'ppe-deploy-phase/). Each phase has its own config.yaml, code, and weights. Baseline outputs are '
    'read-only references; comparison scripts load metrics from cached JSON files, never by importing '
    'baseline code.',
    'Three independent, self-contained project phases. Each is runnable and comparable without '
    'side effects on the others.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 6. Dashboard UI
# ═══════════════════════════════════════════════════════════════════════════════
heading('6. Dashboard UI Iterations', level=1)

heading('6.1 Unclear Model Selection UX', level=2)
challenge_solution(
    'Users saw a "Select Video" dropdown and thought they needed to choose between Baseline '
    'and SAM models before processing.',
    'The sidebar dropdown was originally for switching between pre-computed videos, but its label '
    'and placement made it look like a model selector. Both models always run simultaneously.',
    'Removed the pre-computed video selector entirely since live GPU inference replaced it. '
    'Added clear help text: "Both models run simultaneously, no model selection needed." '
    'Sidebar now shows training metrics and deployment info only.',
    'Cleaner, confusion-free interface. Users upload one video and get both comparisons automatically.'
)

heading('6.2 Results Persistence & Layout', level=2)
challenge_solution(
    'GPU inference results disappeared when the user interacted with other parts of the dashboard. '
    'The pre-computed section cluttered the page below live results.',
    'Streamlit re-runs the entire script on every widget interaction, clearing in-memory state. '
    'Results were stored in a local variable that reset with each re-render.',
    'Stored GPU results in st.session_state.gpu_result with a video-name check. Results persist '
    'across re-renders as long as the same video is selected. Cleaned layout: single scrollable '
    'page with logical flow (KPIs, Compliance, Charts, Preview, Downloads).',
    'Results remain visible throughout user interaction. Page layout flows naturally from '
    'summary to detail.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. Summary
# ═══════════════════════════════════════════════════════════════════════════════
heading('7. Summary: Key Takeaways', level=1)

body('This project exposed challenges across the entire ML deployment stack — from model selection '
     'through cloud infrastructure to UI design. The solutions follow several recurring principles:')

takeaways = [
    ('Prefer stability over novelty. ',
     'SAM 1 over SAM 2, Grounding DINO over Florence-2. The older, battle-tested library almost '
     'always works more reliably than the newer one with exciting features but immature packaging.'),
    ('Test cross-platform early. ',
     'Windows to Linux file transfer (zip paths, line endings) caused multiple failures. A single '
     'test transfer would have caught them before the Colab session.'),
    ('Free tier limitations are real and hard. ',
     'No free persistent GPU hosting exists in 2026. The two-tier architecture (free CPU dashboard '
     '+ serverless GPU) is the pragmatic path forward.'),
    ('Build error isolation into the pipeline. ',
     'Per-model try/except, per-frame fallback, null-safe rendering. Each layer of defense prevents '
     'a single failure from cascading.'),
    ('Large files do not belong in git. ',
     'Video files ballooned the repository to 373 MB, causing Streamlit Cloud clone timeouts. '
     'Clean repos with only code and small data files deploy instantly.'),
    ('Iterate on UX continuously. ',
     'A dropdown labeled "Select Video" confused users into thinking they needed to pick a model. '
     'Small UI fixes have outsized impact on perceived quality.'),
]

for bold, normal in takeaways:
    bold_body(bold, normal)

doc.add_paragraph()
body('Final metrics: mAP50 improved from 0.558 (Baseline) to 0.864 (SAM-Teacher) — a 55% improvement '
     'driven by 9,994 pixel-perfect SAM pseudo-labels. Both models deploy at identical cost, identical '
     'speed (~85 FPS on T4 GPU), and identical model size (40.5 MB). The SAM-Teacher approach delivers '
     'substantially better accuracy with zero runtime overhead.')

# ── Save ──────────────────────────────────────────────────────────────
out_path = r'C:\Users\adm\.openclaw-autoclaw\workspace\ppe-deploy-clean\CHALLENGES_AND_SOLUTIONS.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
